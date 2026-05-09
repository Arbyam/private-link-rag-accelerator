#!/usr/bin/env python3
"""
Sample document generator for Private RAG Accelerator demo.
Generates 5 documents for shared-corpus ingestion.
"""

import os
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def generate_pdf():
    """Generate multi-page PDF with text pages and scanned (OCR) page."""
    pdf_path = os.path.join(OUTPUT_DIR, "private-link-overview.pdf")
    
    doc = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    
    # Page 1: Title and overview
    doc.setFont("Helvetica-Bold", 20)
    doc.drawString(50, height - 50, "Azure Private Link Overview")
    doc.setFont("Helvetica", 12)
    y = height - 100
    
    overview_text = [
        "Azure Private Link provides private connectivity to Azure platform services such as",
        "Azure Storage, Azure SQL Database, and Azure Cosmos DB. It enables you to access",
        "Azure services over a private endpoint in your virtual network.",
        "",
        "Key Benefits:",
        "• Secure connectivity without exposure to public internet",
        "• Consistent network latency for predictable performance",
        "• Compliance with data residency requirements",
        "• Simplified network architecture and reduced attack surface",
        "",
        "Private Link works by routing service traffic through Microsoft's backbone network",
        "rather than traversing the public internet. This ensures your data remains on",
        "Microsoft's private network at all times.",
    ]
    
    for line in overview_text:
        doc.drawString(50, y, line)
        y -= 15
    
    doc.showPage()
    
    # Page 2: Technical details
    doc.setFont("Helvetica-Bold", 20)
    doc.drawString(50, height - 50, "Technical Architecture")
    doc.setFont("Helvetica", 12)
    y = height - 100
    
    technical_text = [
        "Private Link leverages Azure's network virtualization to create private endpoints",
        "within customer virtual networks. A private endpoint is a network interface that",
        "connects you privately and securely to a service powered by Azure Private Link.",
        "",
        "Architecture Components:",
        "1. Private Endpoint: Network interface in your VNet with a private IP address",
        "2. Private Link Resource: The source Azure service (e.g., Storage account)",
        "3. Private DNS Zone: Maintains DNS records for private endpoint resolution",
        "4. Virtual Network: Customer's private network where endpoint resides",
        "",
        "Traffic Flow:",
        "When you access the service via the private endpoint, traffic flows from your",
        "application through the private endpoint directly to the Azure service,",
        "never traversing the public internet. This is achieved through network",
        "virtualization and private DNS integration.",
    ]
    
    for line in technical_text:
        doc.drawString(50, y, line)
        y -= 15
    
    doc.showPage()
    
    # Page 3: Scanned page (OCR exercise)
    # Create an image with text to simulate a scanned document
    img_width, img_height = 800, 1000
    img = Image.new('RGB', (img_width, img_height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Draw a header
    draw.rectangle([0, 0, img_width, 60], fill='lightgray')
    draw.text((20, 10), "CONFIDENTIAL - Deployment Checklist", fill='black')
    
    # Draw scanned-style text (with slight imperfections to look like OCR)
    y_pos = 80
    checklist_text = [
        "Pre-Deployment Checklist for Private Link Implementation",
        "",
        "Network Planning:",
        "[ ] Define IP address space for private endpoints",
        "[ ] Design subnet layout within virtual network",
        "[ ] Plan DNS resolution strategy",
        "[ ] Document firewall rules for NSGs",
        "",
        "Azure Setup:",
        "[ ] Create or identify target Azure services",
        "[ ] Verify service availability in target regions",
        "[ ] Confirm subscription quota limits",
        "[ ] Set up proper RBAC roles",
        "",
        "Security & Compliance:",
        "[ ] Enable network policies on private endpoints",
        "[ ] Configure NSG rules for ingress/egress",
        "[ ] Implement DDoS Protection if needed",
        "[ ] Document data residency requirements",
        "",
        "Monitoring & Management:",
        "[ ] Set up Azure Monitor alerts",
        "[ ] Configure Network Watcher diagnostics",
        "[ ] Plan backup strategy",
        "[ ] Document runbook for failover procedures",
    ]
    
    for line in checklist_text:
        draw.text((30, y_pos), line, fill='black')
        y_pos += 25
    
    # Save image as bytes
    img_bytes = BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    
    # Add image to PDF
    doc.drawImage(ImageReader(img_bytes), 50, 100, width=500, height=600)
    doc.showPage()
    
    doc.save()
    print(f"✓ Generated {pdf_path}")


def generate_txt():
    """Generate RAG architecture notes as plain text."""
    txt_path = os.path.join(OUTPUT_DIR, "rag-architecture-notes.txt")
    
    content = """RAG Architecture and Retrieval-Augmented Generation Fundamentals

Retrieval-Augmented Generation (RAG) is a modern approach to building AI systems that combines
the power of large language models (LLMs) with external knowledge retrieval. Unlike traditional
LLMs that rely solely on their training data, RAG systems dynamically fetch relevant information
from a knowledge base or corpus before generating responses. This approach significantly improves
accuracy, reduces hallucinations, and enables more current information to be incorporated into
model outputs.

Core Components of RAG Architecture:

1. Document Corpus
The foundation of any RAG system is a well-organized collection of documents or data sources.
This corpus contains domain-specific information that provides context for the LLM. In enterprise
settings, this might include internal documentation, product specifications, customer data,
regulatory documents, or research papers. The quality and relevance of the corpus directly
impacts the quality of RAG outputs. Documents must be indexed and stored in a way that enables
efficient retrieval based on semantic similarity or keyword matching.

2. Chunking Strategy
Before documents can be effectively used in RAG, they must be broken down into smaller,
meaningful pieces called chunks. This process, known as chunking, is critical for several reasons:
First, LLM context windows are limited, so large documents must be segmented. Second, chunking
helps ensure that retrieved information is focused and relevant rather than containing extraneous
content. Third, proper chunking preserves semantic meaning within each chunk. Common chunking
strategies include fixed-size chunks with overlap, semantic chunking based on document structure,
or hierarchical chunking that respects document hierarchy. The chunk size typically ranges from
100 to 2000 tokens depending on the use case and model capabilities.

3. Embedding and Vector Storage
To enable semantic search and retrieval, documents are converted into numerical vectors called
embeddings. These embeddings capture the semantic meaning of text, allowing similar concepts to
be represented as nearby points in high-dimensional vector space. Modern embedding models like
OpenAI's text-embedding-3 or Azure's embedding services can convert entire chunks into fixed-size
vectors (typically 768 to 3072 dimensions). These embeddings are stored in a vector database or
search index that supports similarity search. This enables the system to quickly find the most
relevant chunks when a user query arrives, even if the query uses different wording than the
original documents.

4. Retrieval System
The retrieval system is responsible for finding the most relevant documents or chunks given a user
query. It typically works by: (1) converting the query into an embedding using the same model used
for documents, (2) searching the vector index for embeddings most similar to the query embedding,
and (3) returning the top-k most similar chunks. Retrieval can be purely semantic (vector similarity)
or hybrid (combining keyword matching with semantic similarity). Azure AI Search, for example,
supports both approaches. The retrieval system must be fast enough to provide real-time responses,
which is why efficient indexing and search algorithms are essential.

5. Context Window and Prompt Engineering
Once relevant chunks are retrieved, they must be integrated into the prompt sent to the LLM. The
retrieved context is typically inserted into a carefully engineered prompt that instructs the model
to generate a response based on the provided context. Prompt engineering is crucial here—clear
instructions about how to use the retrieved context, what to do if information is not available,
and how to format the response all influence output quality. The context window of the LLM limits
how much retrieved information can be included, so careful selection and summarization may be needed.

6. Response Generation
With the retrieved context included in the prompt, the LLM generates a response. The presence of
relevant context typically results in more accurate, factual, and current responses compared to
using the model's training data alone. The model can now ground its response in the provided
documents, reducing hallucinations and providing verifiable information. It's common to include
citations or references to the source documents in the generated response, enabling users to
verify information and explore source materials.

7. Feedback and Iteration
High-quality RAG systems include mechanisms for collecting feedback on retrieval and generation
quality. This feedback loop enables continuous improvement through: refinement of chunking
strategies, optimization of retrieval parameters, improvement of prompt templates, and retraining
or updating of embedding models as the corpus evolves. Metrics like retrieval precision/recall and
response relevance can guide these improvements.

Implementation Considerations:

Scalability: The system must handle large document collections efficiently. Vector databases need
to support fast similarity search at scale. Index maintenance and update operations must be performant.

Latency: End-to-end latency includes retrieval time plus LLM generation time. Both must be optimized
to provide interactive performance. Caching strategies and efficient inference engines are essential.

Quality Control: Not all retrieved chunks are equally relevant. Filtering by relevance score,
re-ranking retrieved results, or using multiple retrieval passes can improve result quality.

Domain Specificity: Fine-tuning embedding models or retrieval systems for specific domains often
outperforms general-purpose models. Domain-specific training data can significantly improve relevance.

Security and Privacy: When implementing RAG with sensitive documents, access controls, encryption,
and data residency requirements must be carefully managed. Retrieved context should not be logged
or exposed inappropriately.

The Future of RAG:

RAG systems are rapidly evolving with advances in multi-modal embeddings that work with images,
audio, and text; improved retrieval techniques like hybrid search and reranking; and tighter
integration between retrieval and generation components. As these technologies mature, RAG
represents a path toward more accurate, grounded, and controllable AI systems that can leverage
the full breadth of human knowledge while avoiding the hallucinations and latency issues of
purely generative models.
"""
    
    with open(txt_path, 'w') as f:
        f.write(content)
    
    print(f"✓ Generated {txt_path}")


def generate_docx():
    """Generate DOCX document with AI Search deployment guidance."""
    docx_path = os.path.join(OUTPUT_DIR, "ai-search-deployment-guide.docx")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Azure AI Search Deployment Guide', 0)
    
    # Introduction
    doc.add_heading('Introduction', 1)
    intro = doc.add_paragraph(
        'Azure AI Search (formerly Azure Cognitive Search) is a cloud-based search service that '
        'enables developers to build rich search experiences in their applications. This guide '
        'provides best practices and step-by-step instructions for deploying and configuring '
        'AI Search in a production environment with a focus on Private Link connectivity for '
        'enhanced security.'
    )
    
    # Deployment Overview
    doc.add_heading('Deployment Overview', 1)
    deployment_para = doc.add_paragraph(
        'When deploying Azure AI Search, you should consider several key factors: (1) Choosing '
        'the appropriate tier based on your workload requirements (Free, Basic, Standard, etc.), '
        '(2) Selecting the right region for your search service to minimize latency, (3) Configuring '
        'authentication and authorization through Azure AD or key-based authentication, and '
        '(4) Setting up monitoring and alerting for search service health and performance. '
        'For organizations handling sensitive data, integrating Private Link ensures all traffic '
        'between your applications and the search service remains on Microsoft\'s private backbone '
        'network rather than traversing the public internet.'
    )
    
    # Best Practices Section
    doc.add_heading('Best Practices', 1)
    
    practices = [
        'Use managed identities for authentication instead of connection strings to improve security',
        'Implement proper indexing strategies that match your query patterns',
        'Monitor query performance and adjust relevance tuning parameters based on metrics',
        'Enable semantic search and ranking to improve result quality',
        'Configure Network Security Groups to restrict access to the search service endpoint',
        'Use Azure Private Link to establish private connectivity from your applications',
        'Implement query parsing and validation to prevent injection attacks',
        'Set up proper backup and disaster recovery procedures',
        'Use resource locking to prevent accidental modifications to production services'
    ]
    
    for practice in practices:
        doc.add_paragraph(practice, style='List Bullet')
    
    # Configuration
    doc.add_heading('Configuration Steps', 1)
    config_para = doc.add_paragraph(
        'Begin by creating an Azure AI Search service through the Azure portal or using Infrastructure '
        'as Code tools like Terraform or Bicep. Configure the tier based on your document volume and '
        'query throughput requirements. Create a private endpoint and DNS zone to establish secure, '
        'private connectivity. Set up indexers to automatically crawl and index your data sources, '
        'and configure enrichment pipelines to apply AI skills like entity recognition and key phrase '
        'extraction. Finally, configure client applications to connect through the private endpoint '
        'using connection strings or managed identities. Test end-to-end connectivity and perform '
        'load testing to ensure the service meets your performance requirements.'
    )
    
    # Monitoring
    doc.add_heading('Monitoring and Troubleshooting', 1)
    monitor_para = doc.add_paragraph(
        'After deployment, continuously monitor the search service using Azure Monitor. Track key '
        'metrics including query latency, indexing duration, storage utilization, and throttling events. '
        'Configure alerts for anomalous conditions such as unusual latency spikes or indexing errors. '
        'Enable diagnostic logging to capture detailed information about search operations. When issues '
        'occur, check the activity log for service-level events, review indexer execution history, and '
        'validate that network connectivity through private endpoints is functioning correctly. Use the '
        'Search explorer in the portal to test queries and debug relevance issues.'
    )
    
    # Save document
    doc.save(docx_path)
    print(f"✓ Generated {docx_path}")


def generate_png():
    """Generate PNG with network diagram and labels."""
    png_path = os.path.join(OUTPUT_DIR, "network-diagram-with-labels.png")
    
    img_width, img_height = 1000, 800
    img = Image.new('RGB', (img_width, img_height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Background
    draw.rectangle([0, 0, img_width, img_height], outline='black', width=2)
    
    # Title
    draw.text((20, 20), "Azure Private Link Network Architecture", fill='black')
    
    # Draw VNet box
    draw.rectangle([50, 80, 950, 750], outline='darkblue', width=2)
    draw.text((60, 85), "Virtual Network (VNet): 10.0.0.0/16", fill='darkblue')
    
    # Subnet 1
    draw.rectangle([80, 120, 450, 350], outline='blue', width=2)
    draw.text((90, 125), "Subnet 1: 10.0.1.0/24", fill='blue')
    draw.rectangle([100, 150, 250, 200], outline='green', width=2)
    draw.text((110, 160), "App Server", fill='green')
    draw.text((100, 220), "Private IP: 10.0.1.10", fill='black')
    
    # Subnet 2
    draw.rectangle([80, 370, 450, 600], outline='blue', width=2)
    draw.text((90, 375), "Subnet 2: 10.0.2.0/24", fill='blue')
    draw.rectangle([100, 400, 250, 450], outline='purple', width=2)
    draw.text((105, 415), "Database", fill='purple')
    draw.text((100, 470), "Private IP: 10.0.2.20", fill='black')
    
    # Private Endpoint
    draw.rectangle([500, 150, 700, 200], outline='red', width=2)
    draw.text((510, 160), "Private Endpoint", fill='red')
    draw.text((500, 220), "IP: 10.0.3.50", fill='black')
    
    # Azure Service (outside VNet, but connected via PE)
    draw.rectangle([750, 150, 920, 200], outline='orange', width=2)
    draw.text((760, 160), "Azure Storage", fill='orange')
    
    # Connection line from PE to Storage
    draw.line([(700, 175), (750, 175)], fill='red', width=2)
    draw.text((710, 180), "Private Link", fill='red')
    
    # NSG box
    draw.rectangle([500, 370, 920, 600], outline='darkgreen', width=2)
    draw.text((510, 375), "Network Security Group (NSG) Rules", fill='darkgreen')
    draw.text((520, 410), "• Inbound: Allow 10.0.0.0/16 on port 443", fill='black')
    draw.text((520, 440), "• Inbound: Allow 10.0.0.0/16 on port 1433", fill='black')
    draw.text((520, 470), "• Outbound: Allow all to Private Endpoint", fill='black')
    draw.text((520, 500), "• Outbound: Deny internet (except auth)", fill='black')
    
    # Legend
    draw.text((60, 650), "Legend:", fill='black')
    draw.rectangle([60, 670, 100, 690], outline='green', width=2)
    draw.text((110, 672), "Compute", fill='black')
    
    draw.rectangle([250, 670, 290, 690], outline='purple', width=2)
    draw.text((300, 672), "Data Service", fill='black')
    
    draw.rectangle([480, 670, 520, 690], outline='red', width=2)
    draw.text((530, 672), "Private Link Connection", fill='black')
    
    draw.rectangle([750, 670, 790, 690], outline='orange', width=2)
    draw.text((800, 672), "Azure PaaS Service", fill='black')
    
    img.save(png_path)
    print(f"✓ Generated {png_path}")


def generate_jpg():
    """Generate JPG with cost summary chart."""
    jpg_path = os.path.join(OUTPUT_DIR, "cost-summary-chart.jpg")
    
    img_width, img_height = 1000, 800
    img = Image.new('RGB', (img_width, img_height), color='lightyellow')
    draw = ImageDraw.Draw(img)
    
    # Background
    draw.rectangle([0, 0, img_width, img_height], outline='black', width=2, fill='lightyellow')
    
    # Title
    draw.text((20, 20), "Azure Service Monthly Cost Summary", fill='darkred')
    
    # Draw chart border
    draw.rectangle([60, 80, 950, 700], outline='black', width=2)
    
    # Y-axis labels
    draw.text((20, 680), "$0K", fill='black')
    draw.text((20, 580), "$2K", fill='black')
    draw.text((20, 480), "$4K", fill='black')
    draw.text((20, 380), "$6K", fill='black')
    draw.text((20, 280), "$8K", fill='black')
    draw.text((20, 180), "$10K", fill='black')
    
    # Grid lines
    for y in range(80, 700, 100):
        draw.line([(60, y), (950, y)], fill='lightgray', width=1)
    
    # Bar chart data: Service, Cost, Color
    services = [
        ("VMs", 3500, "steelblue"),
        ("Storage", 1200, "forestgreen"),
        ("Private Link", 300, "red"),
        ("AI Search", 2800, "purple"),
        ("Database", 1800, "orange"),
        ("Networking", 400, "darkred"),
    ]
    
    bar_width = 100
    x_start = 100
    max_cost = 10000
    
    for idx, (service, cost, color) in enumerate(services):
        x_pos = x_start + idx * 130
        bar_height = (cost / max_cost) * 600
        
        # Draw bar
        y_top = 680 - bar_height
        draw.rectangle([x_pos, y_top, x_pos + bar_width, 680], fill=color, outline='black', width=1)
        
        # Draw cost label on bar
        cost_text = f"${cost/1000:.1f}K"
        draw.text((x_pos + 10, y_top - 20), cost_text, fill='black')
        
        # Draw service label
        draw.text((x_pos, 710), service, fill='black')
    
    # Legend
    draw.text((60, 750), "Total Monthly Cost: $10,000 | Private Link adds only $300/month for secure connectivity", fill='darkgreen')
    
    # Cost breakdown box
    breakdown_y = 120
    breakdown_lines = [
        "Cost Breakdown:",
        "• Compute (VMs): $3,500 (35%)",
        "• AI Search: $2,800 (28%)",
        "• Database: $1,800 (18%)",
        "• Storage: $1,200 (12%)",
        "• Networking & Private Link: $700 (7%)",
    ]
    
    for line in breakdown_lines:
        draw.text((700, breakdown_y), line, fill='black')
        breakdown_y += 25
    
    img.save(jpg_path, quality=85)
    print(f"✓ Generated {jpg_path}")


def generate_readme():
    """Generate README explaining sample files."""
    readme_path = os.path.join(OUTPUT_DIR, "README.md")
    
    content = """# Sample Documents for Private RAG Accelerator Demo

This directory contains 5 sample documents used to populate the `shared-corpus` container during postprovisioning. These documents are ingested and indexed by the Azure AI Search service, enabling semantic search and retrieval-augmented generation (RAG) in the demo application.

## Files

### 1. private-link-overview.pdf
- **Type**: Multi-page PDF
- **Pages**: 3 (2 text pages + 1 OCR/scanned page)
- **Content**: Azure Private Link concepts, benefits, and technical architecture
- **Purpose**: Exercises text extraction via PyPDF2 and OCR via Document Intelligence on scanned content
- **Size**: ~50 KB

### 2. rag-architecture-notes.txt
- **Type**: Plain text file
- **Content**: Comprehensive guide to Retrieval-Augmented Generation fundamentals
- **Topics**: Document corpus, chunking, embeddings, vector storage, retrieval systems, prompt engineering
- **Purpose**: Provides domain knowledge about RAG for semantic search queries
- **Size**: ~8 KB

### 3. ai-search-deployment-guide.docx
- **Type**: Microsoft Word document
- **Content**: Step-by-step Azure AI Search deployment and configuration guide
- **Sections**: Introduction, deployment overview, best practices, configuration, monitoring
- **Purpose**: Tests DOCX extraction and provides operational guidance
- **Size**: ~15 KB

### 4. network-diagram-with-labels.png
- **Type**: PNG image with rendered text labels
- **Content**: Azure network architecture diagram showing VNet, subnets, private endpoints, NSGs
- **Dimensions**: 1000x800 pixels
- **Purpose**: Exercises OCR on image text and provides visual reference for network design
- **Size**: ~50 KB

### 5. cost-summary-chart.jpg
- **Type**: JPEG image with rendered text and chart
- **Content**: Azure service cost breakdown showing monthly costs by service type
- **Dimensions**: 1000x800 pixels
- **Purpose**: Demonstrates cost analysis and OCR on chart-style images
- **Size**: ~40 KB

## Generation

All files are generated deterministically using a Python script. To regenerate:

```bash
cd samples/
python generate.py
```

### Dependencies

The generation script requires:
- `reportlab` - PDF generation with text rendering
- `python-docx` - DOCX document creation
- `Pillow` - Image creation and manipulation

Install with:
```bash
pip install reportlab python-docx Pillow
```

## Content Relevance

All documents contain real, substantive content about Azure Private Link, RAG architecture, and AI Search deployment. This ensures the demo can:
- Perform meaningful semantic search (queries about "Azure Private Link" or "RAG architecture" will retrieve relevant documents)
- Demonstrate chunking and embedding across multiple document types
- Show accurate citation of sources in RAG responses
- Validate OCR pipeline for scanned content

## Integration

These files are automatically picked up by `scripts/postprovision.ps1` which:
1. Discovers files in the `samples/` directory
2. Filters by allowed MIME types (PDF, TXT, DOCX, PNG, JPEG)
3. Uploads up to 5 files to the `shared-corpus` Azure Blob Storage container
4. Triggers indexing pipeline to process and chunk documents for semantic search

See `scripts/postprovision.ps1` lines 186–246 for implementation details.
"""
    
    with open(readme_path, 'w') as f:
        f.write(content)
    
    print(f"✓ Generated {readme_path}")


if __name__ == "__main__":
    print("Generating sample documents for Private RAG Accelerator...\n")
    generate_pdf()
    generate_txt()
    generate_docx()
    generate_png()
    generate_jpg()
    generate_readme()
    print("\n✓ All samples generated successfully!")
